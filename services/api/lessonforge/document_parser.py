from __future__ import annotations

import hashlib
import io
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED_MEDIA_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
    "text/markdown": ".md",
    "text/x-markdown": ".md",
}


class DocumentParseError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedSegment:
    text: str
    page_number: int | None = None
    paragraph_number: int | None = None


@dataclass(frozen=True)
class ParsedDocument:
    media_type: str
    segments: list[ParsedSegment]

    @property
    def text(self) -> str:
        return "\n\n".join(segment.text for segment in self.segments)


def sanitize_display_name(name: str) -> str:
    cleaned = Path(name.replace("\\", "/")).name
    cleaned = re.sub(r"[\x00-\x1f\x7f]", "", cleaned).strip()
    return cleaned[:255] or "未命名教材"


def detect_media_type(data: bytes, declared_media_type: str, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if data.startswith(b"%PDF-"):
        detected = "application/pdf"
    elif data.startswith(b"PK\x03\x04"):
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                if "word/document.xml" not in archive.namelist():
                    raise DocumentParseError("壓縮檔不是有效的 DOCX 文件")
            detected = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        except zipfile.BadZipFile as error:
            raise DocumentParseError("DOCX 檔案已損壞") from error
    else:
        try:
            data.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise DocumentParseError("僅支援含文字層的 PDF、DOCX、UTF-8 TXT 或 Markdown") from error
        detected = "text/markdown" if suffix in {".md", ".markdown"} else "text/plain"

    allowed_suffixes = {
        "application/pdf": {".pdf"},
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {".docx"},
        "text/plain": {".txt"},
        "text/markdown": {".md", ".markdown"},
    }
    if suffix not in allowed_suffixes[detected]:
        raise DocumentParseError("檔案內容與副檔名不一致")
    normalized_declared = declared_media_type.split(";", maxsplit=1)[0].strip().lower()
    compatible_declared = normalized_declared in {detected, "application/octet-stream"}
    if detected == "text/markdown" and normalized_declared in {"text/plain", "text/x-markdown"}:
        compatible_declared = True
    if not compatible_declared:
        raise DocumentParseError("檔案內容與 MIME type 不一致")
    return detected


def parse_document(data: bytes, media_type: str) -> ParsedDocument:
    if media_type == "application/pdf":
        reader = PdfReader(io.BytesIO(data))
        segments = [
            ParsedSegment(text=text, page_number=index)
            for index, page in enumerate(reader.pages, start=1)
            if (text := (page.extract_text() or "").strip())
        ]
        if not segments:
            raise DocumentParseError("此 PDF 沒有可抽取的文字層；目前版本尚未支援 OCR")
        return ParsedDocument(media_type=media_type, segments=segments)

    if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        document = Document(io.BytesIO(data))
        segments = [
            ParsedSegment(text=paragraph.text.strip(), paragraph_number=index)
            for index, paragraph in enumerate(document.paragraphs, start=1)
            if paragraph.text.strip()
        ]
        if not segments:
            raise DocumentParseError("DOCX 中沒有可抽取的文字")
        return ParsedDocument(media_type=media_type, segments=segments)

    text = data.decode("utf-8-sig").replace("\r\n", "\n").strip()
    if not text:
        raise DocumentParseError("文字檔內容為空")
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    return ParsedDocument(
        media_type=media_type,
        segments=[
            ParsedSegment(text=part, paragraph_number=index)
            for index, part in enumerate(paragraphs, 1)
        ],
    )


def chunk_segments(
    segments: list[ParsedSegment], max_chars: int = 1200, overlap_chars: int = 150
) -> list[ParsedSegment]:
    chunks: list[ParsedSegment] = []
    for segment in segments:
        normalized = re.sub(r"[ \t]+", " ", segment.text).strip()
        if len(normalized) <= max_chars:
            chunks.append(ParsedSegment(normalized, segment.page_number, segment.paragraph_number))
            continue
        start = 0
        while start < len(normalized):
            end = min(start + max_chars, len(normalized))
            if end < len(normalized):
                boundary = normalized.rfind(" ", start, end)
                if boundary > start + max_chars // 2:
                    end = boundary
            chunks.append(
                ParsedSegment(
                    normalized[start:end].strip(), segment.page_number, segment.paragraph_number
                )
            )
            if end >= len(normalized):
                break
            start = max(end - overlap_chars, start + 1)
    return chunks


def sha256_hex(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()
