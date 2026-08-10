from __future__ import annotations

import json
import shutil
import uuid
from io import BytesIO
from pathlib import Path
from typing import Literal

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from jinja2 import Environment, FileSystemLoader, select_autoescape
from playwright.async_api import Error as PlaywrightError
from playwright.async_api import async_playwright

from .config import Settings
from .schemas import PackageView
from .validators import validate_export_content

ExportVariant = Literal["student", "teacher", "homework", "quiz", "parent"]
ExportFormat = Literal["pdf", "docx"]
TEMPLATE_ROOT = Path(__file__).resolve().parents[3] / "templates"


class ExportError(RuntimeError):
    pass


def render_html(
    package: PackageView, *, organization_name: str, class_name: str, variant: ExportVariant
) -> str:
    environment = Environment(
        loader=FileSystemLoader(str(TEMPLATE_ROOT / "pdf")),
        autoescape=select_autoescape(["html", "xml"]),
        trim_blocks=True,
        lstrip_blocks=True,
    )
    html = environment.get_template("lesson.html").render(
        package=package,
        organization_name=organization_name,
        class_name=class_name,
        variant=variant,
        show_answers=variant == "teacher",
    )
    issues = validate_export_content(variant=variant, html=html)
    fatal = [issue for issue in issues if issue.severity == "fatal"]
    if fatal:
        raise ExportError(fatal[0].message)
    return html


def chromium_candidates() -> list[str]:
    paths = [
        shutil.which("chromium"),
        shutil.which("chromium-browser"),
        shutil.which("google-chrome"),
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    ]
    return [path for path in paths if path and Path(path).exists()]


async def html_to_pdf(html: str, destination: Path) -> None:
    async with async_playwright() as playwright:
        browser = None
        errors: list[str] = []
        candidates: list[str | None] = [None, *chromium_candidates()]
        for executable in candidates:
            try:
                browser = await playwright.chromium.launch(
                    headless=True, executable_path=executable
                )
                break
            except PlaywrightError as error:
                errors.append(str(error).splitlines()[0])
        if browser is None:
            raise ExportError(
                "找不到可用的 Chromium／Edge，請執行 `python -m playwright install chromium`。"
                + " | ".join(errors[:2])
            )
        try:
            page = await browser.new_page()
            await page.set_content(html, wait_until="networkidle")
            await page.emulate_media(media="print")
            await page.pdf(
                path=str(destination),
                format="A4",
                print_background=True,
                margin={"top": "16mm", "right": "14mm", "bottom": "18mm", "left": "14mm"},
                display_header_footer=False,
            )
        finally:
            await browser.close()


def set_east_asian_font(run: object, name: str) -> None:
    run.font.name = name  # type: ignore[attr-defined]
    run._element.rPr.rFonts.set(qn("w:ascii"), name)  # type: ignore[attr-defined]
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)  # type: ignore[attr-defined]
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)  # type: ignore[attr-defined]


def set_style_font(style: object, name: str, size: float, color: RGBColor) -> None:
    style.font.name = name  # type: ignore[attr-defined]
    style.font.size = Pt(size)  # type: ignore[attr-defined]
    style.font.color.rgb = color  # type: ignore[attr-defined]
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)  # type: ignore[attr-defined]
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)  # type: ignore[attr-defined]
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)  # type: ignore[attr-defined]


def keep_with_next(paragraph: object) -> None:
    paragraph.paragraph_format.keep_with_next = True  # type: ignore[attr-defined]


def add_section_title(
    document: DocxDocument, text: str, *, font_name: str, accent: RGBColor
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_east_asian_font(run, font_name)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = accent


def source_label(reference: object) -> str:
    if reference.page_number:  # type: ignore[attr-defined]
        location = f"第 {reference.page_number} 頁"  # type: ignore[attr-defined]
    else:
        location = f"段落 {reference.paragraph_number or '-'}"  # type: ignore[attr-defined]
    return f"{reference.material_name}（{location}）"  # type: ignore[attr-defined]


def add_page_number(section: object) -> None:
    footer = section.footer  # type: ignore[attr-defined]
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)
    paragraph.add_run(" 頁")


def build_docx(
    package: PackageView,
    *,
    organization_name: str,
    class_name: str,
    variant: ExportVariant,
) -> bytes:
    style_settings = json.loads(
        (TEMPLATE_ROOT / "docx" / "styles.json").read_text(encoding="utf-8")
    )
    font_name = style_settings["font"]
    accent = RGBColor.from_string(style_settings["accent"].lstrip("#"))
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    add_page_number(section)
    normal = document.styles["Normal"]
    set_style_font(normal, font_name, 10.5, RGBColor(23, 35, 31))
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    heading_1 = document.styles["Heading 1"]
    set_style_font(heading_1, font_name, 14, accent)
    heading_1.paragraph_format.space_before = Pt(14)
    heading_1.paragraph_format.space_after = Pt(7)
    heading_1.paragraph_format.keep_with_next = True
    heading_2 = document.styles["Heading 2"]
    set_style_font(heading_2, font_name, 12, RGBColor(23, 79, 69))
    heading_2.paragraph_format.space_before = Pt(10)
    heading_2.paragraph_format.space_after = Pt(5)
    heading_2.paragraph_format.keep_with_next = True
    list_bullet = document.styles["List Bullet"]
    set_style_font(list_bullet, font_name, 10.5, RGBColor(23, 35, 31))
    list_bullet.paragraph_format.left_indent = Cm(0.95)
    list_bullet.paragraph_format.first_line_indent = Cm(-0.48)
    list_bullet.paragraph_format.space_after = Pt(4)
    list_bullet.paragraph_format.line_spacing = 1.25

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run(package.title)
    set_east_asian_font(title_run, font_name)
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = accent
    subtitle = document.add_paragraph(
        f"{organization_name}｜{class_name}｜{package.lesson_date.isoformat()}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.keep_with_next = True

    if variant in {"student", "teacher"}:
        for block in package.blocks:
            heading = document.add_heading(
                f"{block.title}（{block.duration_minutes} 分鐘）", level=1
            )
            for run in heading.runs:
                set_east_asian_font(run, font_name)
            if block.student_content:
                document.add_paragraph(block.student_content)
            if block.instructions:
                document.add_paragraph(block.instructions)
            for index, question in enumerate(block.questions, 1):
                question_paragraphs: list[object] = []
                question_paragraphs.append(document.add_paragraph(f"{index}. {question.prompt}"))
                for option_index, option in enumerate(question.options):
                    question_paragraphs.append(
                        document.add_paragraph(
                            f"{chr(65 + option_index)}. {option}", style="List Bullet"
                        )
                    )
                if variant == "student":
                    question_paragraphs.append(
                        document.add_paragraph(
                            "作答：________________________________________________"
                        )
                    )
                else:
                    answer = document.add_paragraph()
                    answer_run = answer.add_run(f"答案：{question.answer}")
                    answer_run.bold = True
                    answer_run.font.color.rgb = accent
                    question_paragraphs.append(answer)
                    question_paragraphs.append(
                        document.add_paragraph(f"解析：{question.explanation}")
                    )
                for paragraph in question_paragraphs[:-1]:
                    keep_with_next(paragraph)
            if variant == "teacher":
                document.add_paragraph(f"教師備註：{block.teacher_notes}")
                if block.source_references:
                    sources = "；".join(source_label(ref) for ref in block.source_references)
                    document.add_paragraph(f"來源：{sources}")
    elif variant == "homework":
        for day in package.homework_days:
            document.add_heading(day.title, level=1)
            document.add_paragraph(
                f"預估 {day.estimated_minutes} 分鐘｜單字：{'、'.join(day.vocabulary)}"
            )
            for index, question in enumerate(day.questions, 1):
                document.add_paragraph(
                    f"{index}. {question.prompt}\n作答：________________________________________"
                )
            document.add_paragraph(day.review_note)
        document.add_section(WD_SECTION.NEW_PAGE)
        add_section_title(document, "作業完整答案", font_name=font_name, accent=accent)
        for day in package.homework_days:
            for index, question in enumerate(day.questions, 1):
                document.add_paragraph(
                    f"Day {day.day}-{index}：{question.answer}｜{question.explanation}"
                )
    elif variant == "quiz" and package.weekly_quiz:
        quiz = package.weekly_quiz
        document.add_paragraph(
            f"總分 {quiz.total_points} 分｜建議作答時間 {quiz.suggested_minutes} 分鐘"
        )
        for index, question in enumerate(quiz.questions, 1):
            document.add_paragraph(f"{index}. {question.prompt}（{question.points} 分）")
            for option_index, option in enumerate(question.options):
                document.add_paragraph(f"{chr(65 + option_index)}. {option}")
        document.add_section(WD_SECTION.NEW_PAGE)
        add_section_title(document, "週考答案卷", font_name=font_name, accent=accent)
        for index, question in enumerate(quiz.questions, 1):
            document.add_paragraph(f"{index}. {question.answer}｜{question.explanation}")
    elif variant == "parent" and package.parent_report:
        report = package.parent_report
        items = [
            ("作業完成情況", report.homework_completion),
            ("本週測驗表現", report.quiz_performance),
            ("本週進步", report.progress),
            ("主要弱點", "、".join(report.main_weaknesses)),
            ("下週教學重點", "、".join(report.next_week_focus)),
            ("教師備註", report.teacher_notes),
        ]
        for label, value in items:
            document.add_heading(label, level=1)
            document.add_paragraph(value)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def export_document(
    package: PackageView,
    *,
    settings: Settings,
    organization_name: str,
    class_name: str,
    variant: ExportVariant,
    file_format: ExportFormat,
) -> tuple[Path, str]:
    settings.ensure_directories()
    destination = settings.export_dir / f"{uuid.uuid4()}.{file_format}"
    if file_format == "pdf":
        html = render_html(
            package, organization_name=organization_name, class_name=class_name, variant=variant
        )
        await html_to_pdf(html, destination)
        return destination, "application/pdf"
    destination.write_bytes(
        build_docx(
            package, organization_name=organization_name, class_name=class_name, variant=variant
        )
    )
    return destination, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
